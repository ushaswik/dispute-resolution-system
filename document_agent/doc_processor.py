"""
Document Processor Agent - Production Architecture
Supports: TXT, PDF, DOCX, XLSX, JPG, PNG
Uses: Tesseract OCR + GPT-4o/GPT-4o-mini
"""

import os
import re
import base64
import json
from typing import Dict, List, Optional, Tuple
from datetime import datetime

# Document processing
import PyPDF2
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
import docx
import openpyxl

import pytesseract

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# AI
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()

# Debug startup marker
print("[doc-processor] startup")


class DocumentProcessor:
    """
    Multi-format document processor with intelligent routing
    """
    
    # Confidence threshold for OCR vs Vision fallback
    OCR_CONFIDENCE_THRESHOLD = 70.0
    
    # Supported file types
    SUPPORTED_EXTENSIONS = {
        'text': ['.txt'],
        'word': ['.docx', '.doc'],
        'excel': ['.xlsx', '.xls'],
        'pdf': ['.pdf'],
        'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']
    }
    
    def __init__(self):
        """Initialize OpenAI client and OCR"""
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY not found")
        
        self.client = OpenAI(api_key=api_key)
        self.text_model = "gpt-4o-mini"  # For text documents
        self.vision_model = "gpt-4o"      # For vision tasks
        
        # Configure Tesseract path if on Windows
        # Uncomment and set path if needed:
        # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    
    def process_document(
        self, 
        file_path: str, 
        dispute_context: Optional[Dict] = None
    ) -> Dict:
        """
        Main entry point - processes any supported document type
        
        Args:
            file_path: Path to document
            dispute_context: Optional context about the dispute
            
        Returns:
            Structured extraction results
        """
        
        # Validate file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Detect file type
        file_type, extension = self._detect_file_type(file_path)
        
        print(f"📄 Processing {file_type.upper()} file: {os.path.basename(file_path)}")
        
        # Route to appropriate processor
        if file_type == 'text':
            result = self._process_text_file(file_path, dispute_context)
            
        elif file_type == 'word':
            result = self._process_word_file(file_path, dispute_context)
            
        elif file_type == 'excel':
            result = self._process_excel_file(file_path, dispute_context)
            
        elif file_type == 'pdf':
            result = self._process_pdf_file(file_path, dispute_context)
            
        elif file_type == 'image':
            result = self._process_image_file(file_path, dispute_context)
            
        else:
            raise ValueError(f"Unsupported file type: {extension}")
        
        # Add metadata
        result['file_name'] = os.path.basename(file_path)
        result['file_type'] = file_type
        result['processing_timestamp'] = datetime.utcnow().isoformat()
        
        return result
    
    def _detect_file_type(self, file_path: str) -> Tuple[str, str]:
        """Detect file type from extension"""
        extension = os.path.splitext(file_path)[1].lower()
        
        for file_type, extensions in self.SUPPORTED_EXTENSIONS.items():
            if extension in extensions:
                return file_type, extension
        
        return 'unknown', extension
    
    # ==================== TEXT FILES ====================
    
    def _process_text_file(self, file_path: str, context: Optional[Dict]) -> Dict:
        """Process plain text files"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        return self._interpret_with_llm(
            text=text,
            context=context,
            method="text_parse"
        )
    
    # ==================== WORD FILES ====================
    
    def _process_word_file(self, file_path: str, context: Optional[Dict]) -> Dict:
        """Process Word documents (.docx)"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            text = "\n".join([para.text for para in doc.paragraphs])
            
            # Extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += "\n" + row_text
            
            return self._interpret_with_llm(
                text=text,
                context=context,
                method="docx_parse"
            )
            
        except Exception as e:
            raise Exception(f"Failed to process Word document: {e}")
    
    # ==================== EXCEL FILES ====================
    
    def _process_excel_file(self, file_path: str, context: Optional[Dict]) -> Dict:
        """Process Excel files (.xlsx)"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            # Get first sheet
            sheet = workbook.active
            
            # Extract all cell values
            text_rows = []
            for row in sheet.iter_rows(values_only=True):
                # Filter out None values and convert to string
                row_text = " | ".join([str(cell) for cell in row if cell is not None])
                if row_text.strip():
                    text_rows.append(row_text)
            
            text = "\n".join(text_rows)
            
            return self._interpret_with_llm(
                text=text,
                context=context,
                method="xlsx_parse"
            )
            
        except Exception as e:
            raise Exception(f"Failed to process Excel file: {e}")
    
    # ==================== PDF FILES ====================
    
    def _process_pdf_file(self, file_path: str, context: Optional[Dict]) -> Dict:
        """
        Process PDF files with fallback strategy:
        1. Try text extraction (PyPDF2)
        2. If fails, use Tesseract OCR
        3. If OCR confidence low, fallback to Vision
        """
        
        # Step 1: Try text extraction
        print("  → Attempting text extraction...")
        text = self._extract_pdf_text(file_path)
        
        if len(text.strip()) > 50:  # Successfully extracted meaningful text
            print("  ✓ Text extraction successful")
            return self._interpret_with_llm(
                text=text,
                context=context,
                method="pdf_text_parse"
            )
        
        # Step 2: PDF is likely scanned - use OCR
        print("  → Text extraction failed, using OCR...")
        ocr_text, confidence = self._ocr_pdf(file_path)
        
        if confidence >= self.OCR_CONFIDENCE_THRESHOLD:
            print(f"  ✓ OCR successful (confidence: {confidence:.1f}%)")
            return self._interpret_with_llm(
                text=ocr_text,
                context=context,
                method="pdf_ocr",
                ocr_confidence=confidence
            )
        
        # Step 3: OCR confidence low - fallback to Vision
        print(f"  → OCR confidence low ({confidence:.1f}%), using Vision...")
        return self._process_pdf_with_vision(file_path, context)
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using PyPDF2"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                return text
        except Exception as e:
            print(f"  ⚠ PDF text extraction error: {e}")
            return ""
    
    def _ocr_pdf(self, file_path: str) -> Tuple[str, float]:
        """
        OCR a PDF file using Tesseract
        Returns: (extracted_text, confidence_score)
        """
        try:
            # Convert PDF to images (one per page)
            images = convert_from_path(file_path, dpi=300)
            
            all_text = []
            all_confidences = []
            
            for i, image in enumerate(images):
                # Get OCR data with confidence
                ocr_data = pytesseract.image_to_data(
                    image, 
                    output_type=pytesseract.Output.DICT
                )
                
                # Extract text
                page_text = " ".join([
                    word for word in ocr_data['text'] 
                    if word.strip()
                ])
                all_text.append(page_text)
                
                # Calculate average confidence for this page
                confidences = [
                    conf for conf in ocr_data['conf'] 
                    if conf != -1  # -1 means no confidence available
                ]
                if confidences:
                    avg_conf = sum(confidences) / len(confidences)
                    all_confidences.append(avg_conf)
            
            # Combine all pages
            full_text = "\n".join(all_text)
            
            # Overall confidence
            overall_confidence = (
                sum(all_confidences) / len(all_confidences) 
                if all_confidences else 0.0
            )
            
            return full_text, overall_confidence
            
        except Exception as e:
            print(f"  ⚠ OCR error: {e}")
            return "", 0.0
    
    def _process_pdf_with_vision(self, file_path: str, context: Optional[Dict]) -> Dict:
        """Process PDF using GPT-4o Vision (converts first page to image)"""
        try:
            # Convert first page to image
            images = convert_from_path(file_path, dpi=300, first_page=1, last_page=1)
            
            if not images:
                raise Exception("Could not convert PDF to image")
            
            # Save temporary image
            temp_image_path = file_path.replace('.pdf', '_temp.jpg')
            images[0].save(temp_image_path, 'JPEG')
            
            # Process with vision
            result = self._process_with_vision(temp_image_path, context)
            result['extraction_method'] = 'pdf_vision'
            
            # Cleanup
            if os.path.exists(temp_image_path):
                os.remove(temp_image_path)
            
            return result
            
        except Exception as e:
            raise Exception(f"Vision processing failed: {e}")
    
    # ==================== IMAGE FILES ====================
    
    def _process_image_file(self, file_path: str, context: Optional[Dict]) -> Dict:
        """
        Process image files:
        1. Try Tesseract OCR
        2. If confidence low, fallback to Vision
        """
        
        print("  → Attempting OCR...")
        ocr_text, confidence = self._ocr_image(file_path)
        
        if confidence >= self.OCR_CONFIDENCE_THRESHOLD:
            print(f"  ✓ OCR successful (confidence: {confidence:.1f}%)")
            return self._interpret_with_llm(
                text=ocr_text,
                context=context,
                method="image_ocr",
                ocr_confidence=confidence
            )
        
        print(f"  → OCR confidence low ({confidence:.1f}%), using Vision...")
        return self._process_with_vision(file_path, context)
    
    def _ocr_image(self, file_path: str) -> Tuple[str, float]:
        """
        OCR an image using Tesseract
        Returns: (extracted_text, confidence_score)
        """
        try:
            image = Image.open(file_path)
            
            # Get OCR data with confidence
            ocr_data = pytesseract.image_to_data(
                image,
                output_type=pytesseract.Output.DICT
            )
            
            # Extract text
            text = " ".join([
                word for word in ocr_data['text']
                if word.strip()
            ])
            
            # Calculate average confidence
            confidences = [
                conf for conf in ocr_data['conf']
                if conf != -1
            ]
            
            avg_confidence = (
                sum(confidences) / len(confidences)
                if confidences else 0.0
            )
            
            return text, avg_confidence
            
        except Exception as e:
            print(f"  ⚠ Image OCR error: {e}")
            return "", 0.0
    
    def _process_with_vision(self, image_path: str, context: Optional[Dict]) -> Dict:
        """Process image using GPT-4o Vision"""
        
        # Encode image to base64
        with open(image_path, "rb") as image_file:
            base64_image = base64.b64encode(image_file.read()).decode('utf-8')
        
        # Determine MIME type
        ext = os.path.splitext(image_path)[1].lower()
        mime_type = f"image/{ext[1:] if ext != '.jpg' else 'jpeg'}"
        
        # Build prompt
        prompt = self._build_extraction_prompt(context)
        
        try:
            response = self.client.chat.completions.create(
                model=self.vision_model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert at extracting information from receipts, invoices, and financial documents. Always respond in valid JSON format."
                    },
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                temperature=0.1,
                max_tokens=1500,
                response_format={"type": "json_object"}
            )
            
            content = response.choices[0].message.content
            try:
                result = json.loads(content)
            except Exception:
                # Fallback: convert JS-style literals to Python and eval safely
                import ast
                content_fixed = (
                    content.replace('null', 'None')
                           .replace('true', 'True')
                           .replace('false', 'False')
                )
                result = ast.literal_eval(content_fixed)
            result['extraction_method'] = 'vision'
            
            return self._validate_result(result)
            
        except Exception as e:
            raise Exception(f"Vision API error: {e}")
    
    # ==================== LLM INTERPRETATION ====================
    
    def _interpret_with_llm(
        self,
        text: str,
        context: Optional[Dict],
        method: str,
        ocr_confidence: Optional[float] = None
    ) -> Dict:
        """Interpret extracted text using GPT"""
        
        prompt = self._build_extraction_prompt(context, text)
        
        try:
            response = self.client.chat.completions.create(
                model=self.text_model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert at extracting structured information from receipts, invoices, and financial documents. Always respond in valid JSON format."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.1,
                response_format={"type": "json_object"}
            )
            
            content = response.choices[0].message.content
            try:
                result = json.loads(content)
            except Exception:
                import ast
                content_fixed = (
                    content.replace('null', 'None')
                           .replace('true', 'True')
                           .replace('false', 'False')
                )
                result = ast.literal_eval(content_fixed)
            result['extraction_method'] = method
            
            if ocr_confidence is not None:
                result['ocr_confidence'] = round(ocr_confidence, 2)
            
            return self._validate_result(result)
            
        except Exception as e:
            raise Exception(f"LLM interpretation error: {e}")
    
    # ==================== PROMPTS ====================
    
    def _build_extraction_prompt(
        self, 
        context: Optional[Dict], 
        text: Optional[str] = None
    ) -> str:
        """Build extraction prompt for LLM"""
        
        context_info = ""
        if context:
            context_info = f"""
Context about this dispute:
- Dispute type: {context.get('intent_name', 'Unknown')}
- Disputed amount: ${context.get('disputed_amount', 'Unknown')}
- Transaction date: {context.get('transaction_date', 'Unknown')}
"""
        
        text_section = f"\n\nDocument text:\n{text}\n" if text else ""
        
        prompt = f"""
Extract information from this document.

{context_info}{text_section}

Return a JSON object with these fields:
{{
    "merchant_name": "name of merchant/store",
    "transaction_date": "date in YYYY-MM-DD format",
    "transaction_amount": numeric total amount,
    "transaction_currency": "USD|EUR|etc",
    "items_purchased": ["list of items if available"],
    "payment_method": "payment method description",
    "card_last_four": "last 4 digits of card if present",
    "order_number": "order/transaction number if present",
    "document_type": "receipt|invoice|statement|letter|spreadsheet|other",
    "key_evidence": ["facts relevant to a payment dispute"],
    "supports_dispute": true or false,
    "dispute_relevance": "high|medium|low",
    "confidence": number 0.0-1.0
}}

Rules:
- Use null for fields not found
- Convert all dates to YYYY-MM-DD
- Extract only final total amount
- For key_evidence, identify timing, amounts, discrepancies
- Assess if document supports the cardholder's claim
"""
        return prompt
    
    # ==================== VALIDATION ====================
    
    def _validate_result(self, result: Dict) -> Dict:
        """Validate and clean extraction result"""
        
        required_fields = [
            "merchant_name", "transaction_date", "transaction_amount",
            "document_type", "confidence"
        ]
        
        # Ensure required fields exist
        for field in required_fields:
            if field not in result:
                result[field] = None
        
        # Validate confidence
        if result.get("confidence") is not None:
            result["confidence"] = max(0.0, min(1.0, float(result["confidence"])))
        else:
            result["confidence"] = 0.5
        
        return result


# ==================== TEST FUNCTION ====================

if __name__ == "__main__":
    processor = DocumentProcessor()
    
    print("=" * 80)
    print("MULTI-FORMAT DOCUMENT PROCESSOR TEST")
    print("=" * 80)
    print()
    
    # Test files
    test_files = [
        "test_documents/sample_receipt.txt",
        "test_documents/receipt1.jpg",
        # Add more test files as you create them
    ]
    
    context = {
        "intent_name": "Lost/Stolen Card",
        "disputed_amount": 12.75,
        "transaction_date": "2026-03-07"
    }
    
    for file_path in test_files:
        if os.path.exists(file_path):
            try:
                print()
                result = processor.process_document(file_path, context)
                
                print("\n✅ EXTRACTION RESULTS:")
                print("-" * 80)
                for key, value in result.items():
                    print(f"  {key}: {value}")
                print("-" * 80)
                
            except Exception as e:
                print(f"\n❌ ERROR: {e}\n")
        else:
            print(f"⚠️  File not found: {file_path}\n")
    
    print("\n" + "=" * 80)