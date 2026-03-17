"""
Create test documents in various formats for testing
"""

import os
from PIL import Image, ImageDraw, ImageFont
import docx
import openpyxl
from openpyxl.styles import Font

# Create test_documents directory next to this script
BASE_DIR = os.path.dirname(__file__)
TARGET_DIR = os.path.join(BASE_DIR, "test_documents")
os.makedirs(TARGET_DIR, exist_ok=True)

# ==================== TEXT FILE ====================
print(f"Creating {os.path.join(TARGET_DIR, 'sample_receipt.txt')}...")
with open(os.path.join(TARGET_DIR, "sample_receipt.txt"), "w") as f:
    f.write("""TARGET STORE #1234
123 Main Street, Seattle, WA 98101
(206) 555-0100

Date: 03/07/2026
Time: 14:35:22
Transaction #: T789456123

ITEMS:
Wireless Mouse          $24.99
USB-C Cable             $12.99
Laptop Stand            $45.99
-------------------------
Subtotal:               $83.97
Tax (10.1%):            $8.48
TOTAL:                  $92.45

Payment: Visa ending in 4567
Approval Code: 456789

Thank you for shopping at Target!
""")
print("✅ Created sample_receipt.txt")

# ==================== WORD DOCUMENT ====================
print(f"Creating {os.path.join(TARGET_DIR, 'invoice.docx')}...")
doc = docx.Document()
doc.add_heading('INVOICE', 0)

doc.add_paragraph('Amazon.com Services LLC')
doc.add_paragraph('410 Terry Ave N, Seattle, WA 98109')
doc.add_paragraph()

doc.add_paragraph(f'Invoice Date: March 7, 2026')
doc.add_paragraph(f'Invoice #: INV-2026-001234')
doc.add_paragraph()

doc.add_heading('Bill To:', level=2)
doc.add_paragraph('John Doe')
doc.add_paragraph('456 Oak Street')
doc.add_paragraph('Portland, OR 97201')
doc.add_paragraph()

doc.add_heading('Items:', level=2)

# Create table
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Description'
hdr_cells[1].text = 'Qty'
hdr_cells[2].text = 'Amount'

# Data rows
row1 = table.rows[1].cells
row1[0].text = 'Echo Dot (4th Gen)'
row1[1].text = '2'
row1[2].text = '$89.98'

row2 = table.rows[2].cells
row2[0].text = 'Fire TV Stick'
row2[1].text = '1'
row2[2].text = '$39.99'

row3 = table.rows[3].cells
row3[0].text = 'TOTAL'
row3[1].text = ''
row3[2].text = '$129.97'

doc.add_paragraph()
doc.add_paragraph('Payment Method: Visa ending in 4567')
doc.add_paragraph('Transaction Date: March 7, 2026')

doc.save(os.path.join(TARGET_DIR, 'invoice.docx'))
print("✅ Created invoice.docx")

# ==================== EXCEL FILE ====================
print(f"Creating {os.path.join(TARGET_DIR, 'expense_report.xlsx')}...")
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Expense Report"

# Headers
ws['A1'] = 'EXPENSE REPORT'
ws['A1'].font = Font(size=14, bold=True)

ws['A3'] = 'Employee:'
ws['B3'] = 'Jane Smith'
ws['A4'] = 'Department:'
ws['B4'] = 'Sales'
ws['A5'] = 'Date:'
ws['B5'] = '03/07/2026'

# Expense table
ws['A7'] = 'Date'
ws['B7'] = 'Merchant'
ws['C7'] = 'Category'
ws['D7'] = 'Amount'

expenses = [
    ['03/05/2026', 'Starbucks', 'Meals', 15.75],
    ['03/06/2026', 'Uber', 'Transportation', 28.50],
    ['03/07/2026', 'Best Western', 'Lodging', 145.00],
    ['03/07/2026', 'Shell Gas', 'Fuel', 52.30],
]

for i, expense in enumerate(expenses, start=8):
    ws[f'A{i}'] = expense[0]
    ws[f'B{i}'] = expense[1]
    ws[f'C{i}'] = expense[2]
    ws[f'D{i}'] = expense[3]

ws['C13'] = 'TOTAL:'
ws['D13'] = '=SUM(D8:D11)'
ws['D13'].font = Font(bold=True)

wb.save(os.path.join(TARGET_DIR, 'expense_report.xlsx'))
print("✅ Created expense_report.xlsx")

# ==================== IMAGE RECEIPT ====================
print(f"Creating {os.path.join(TARGET_DIR, 'receipt_image.jpg')}...")
img = Image.new('RGB', (600, 900), color='white')
draw = ImageDraw.Draw(img)

try:
    font_large = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 28)
    font_medium = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 20)
    font_small = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 16)
except:
    font_large = font_medium = font_small = ImageFont.load_default()

y = 40

# Store header
draw.text((150, y), "WHOLE FOODS MARKET", fill='black', font=font_large)
y += 50
draw.text((180, y), "456 Pine St, Seattle", fill='black', font=font_small)
y += 25
draw.text((220, y), "(206) 555-0200", fill='black', font=font_small)
y += 50

# Date/time
draw.text((50, y), "Date: 03/07/2026", fill='black', font=font_medium)
draw.text((350, y), "Time: 09:15 AM", fill='black', font=font_medium)
y += 40

draw.text((50, y), "Transaction: WF-789456", fill='black', font=font_medium)
y += 60

# Line
draw.line([(40, y), (560, y)], fill='black', width=2)
y += 30

# Items
items = [
    ("Organic Bananas", "2.49"),
    ("Almond Milk", "4.99"),
    ("Whole Grain Bread", "5.49"),
    ("Free Range Eggs", "6.99"),
    ("Organic Spinach", "3.99"),
]

for item, price in items:
    draw.text((60, y), item, fill='black', font=font_small)
    draw.text((460, y), f"${price}", fill='black', font=font_small)
    y += 28

y += 20
draw.line([(40, y), (560, y)], fill='black', width=1)
y += 30

# Totals
draw.text((60, y), "Subtotal:", fill='black', font=font_medium)
draw.text((460, y), "$23.95", fill='black', font=font_medium)
y += 35

draw.text((60, y), "Tax (10.1%):", fill='black', font=font_medium)
draw.text((460, y), "$2.42", fill='black', font=font_medium)
y += 40

draw.text((60, y), "TOTAL:", fill='black', font=font_large)
draw.text((450, y), "$26.37", fill='black', font=font_large)
y += 60

draw.line([(40, y), (560, y)], fill='black', width=2)
y += 30

# Payment
draw.text((50, y), "Payment: Mastercard x-1234", fill='black', font=font_medium)
y += 50

draw.text((150, y), "Thank you for shopping!", fill='black', font=font_medium)

img.save(os.path.join(TARGET_DIR, 'receipt_image.jpg'))
print("✅ Created receipt_image.jpg")

print("\n" + "="*60)
print("All test documents created successfully!")
print("="*60)